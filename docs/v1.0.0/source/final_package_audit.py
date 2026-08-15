from __future__ import annotations

import hashlib
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


QA = Path(__file__).resolve().parent
STAGING = QA / "staging"
SPLIT_MD = QA / "split-staging"
SPLIT_DOCX = QA / "split-docx"
# This file lives at <project>/docs/v1.0.0/source/.
OUTPUT = QA.parent
PROJECT = OUTPUT.parents[1]
SECTION_DOCX = OUTPUT / "sections"
SECTION_MD = OUTPUT / "markdown" / "sections"

NS = {"cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties", "dc": "http://purl.org/dc/elements/1.1/"}


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def text_ids(pattern: str, path: Path) -> set[str]:
    return set(re.findall(pattern, path.read_text(encoding="utf-8")))


def inspect_docx(path: Path) -> list[str]:
    errors: list[str] = []
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        if "word/comments.xml" in names or "word/commentsExtended.xml" in names:
            errors.append(f"{path.name}: comments present")
        main = archive.read("word/document.xml")
        if re.search(br"<w:(?:ins|del)(?:\s|>)", main):
            errors.append(f"{path.name}: tracked changes present")
        if "docProps/core.xml" in names:
            root = ET.fromstring(archive.read("docProps/core.xml"))
            for tag in ("dc:creator", "cp:lastModifiedBy"):
                node = root.find(tag, NS)
                if node is not None and (node.text or "").strip():
                    errors.append(f"{path.name}: non-anonymous {tag}")
    document = Document(path)
    if not document.paragraphs or not document.paragraphs[0].text.strip():
        errors.append(f"{path.name}: title missing")
    if not any(p.style and p.style.name.startswith("Heading") for p in document.paragraphs):
        errors.append(f"{path.name}: no Word heading styles")
    return errors


def main() -> None:
    errors: list[str] = []
    sources = sorted(STAGING.glob("*.md"))
    if len(sources) != 6:
        errors.append(f"expected six source Markdown files, found {len(sources)}")

    all_text = "\n".join(path.read_text(encoding="utf-8") for path in sources)
    for term in ("ADR-", "roadmap", "milestone", "work package", "project governance"):
        if re.search(re.escape(term), all_text, re.IGNORECASE):
            errors.append(f"prohibited term found: {term}")
    if re.search(r"(?i)(?:requires?|uses?|depends? on) (?:an? )?(?:App Lock )?Accessibility (?:service|permission)", all_text):
        errors.append("App Lock Accessibility service appears required")
    for term in ("Android 11", "Android 15", "API levels 30 through 35", "Usage Access"):
        if term not in all_text:
            errors.append(f"scope term missing: {term}")

    source_fr = set()
    for p in (PROJECT / "docs" / "v2.0.0" / "srs" / "md").glob("*.md"):
        source_fr |= text_ids(r"\bFR-\d{3}\b", p)
    source_nfr = set()
    for p in (PROJECT / "docs" / "v2.0.0" / "nfr" / "md").glob("*.md"):
        source_nfr |= text_ids(r"\bNFR-[A-Z]+-\d{3}\b", p)
    srs = STAGING / "SRS_v1_0_0.md"
    nfr = STAGING / "NFR_v1_0_0.md"
    if text_ids(r"\bFR-\d{3}\b", srs) != source_fr:
        errors.append("SRS requirement identifier accounting mismatch")
    if text_ids(r"\bNFR-[A-Z]+-\d{3}\b", nfr) != source_nfr:
        errors.append("NFR requirement identifier accounting mismatch")

    source_mds = sorted(SPLIT_MD.glob("*.md"))
    source_docx = sorted(SPLIT_DOCX.glob("*.docx"))
    if len(source_mds) != 118 or len(source_docx) != 118:
        errors.append(f"staged split count mismatch: Markdown={len(source_mds)}, DOCX={len(source_docx)}")

    SECTION_DOCX.mkdir(parents=True, exist_ok=True)
    SECTION_MD.mkdir(parents=True, exist_ok=True)
    for path in source_mds:
        shutil.copy2(path, SECTION_MD / path.name)
    for path in source_docx:
        shutil.copy2(path, SECTION_DOCX / path.name)

    final_mds = sorted(SECTION_MD.glob("*.md"))
    final_docx = sorted(SECTION_DOCX.glob("*.docx"))
    if {p.stem for p in final_mds} != {p.stem for p in final_docx}:
        errors.append("final section Markdown/DOCX basename sets differ")
    for path in source_mds:
        target = SECTION_MD / path.name
        if not target.exists() or digest(path) != digest(target):
            errors.append(f"Markdown mirror copy mismatch: {path.name}")
    for path in source_docx:
        target = SECTION_DOCX / path.name
        if not target.exists() or digest(path) != digest(target):
            errors.append(f"DOCX copy mismatch: {path.name}")

    all_docx = sorted(OUTPUT.glob("*.docx")) + final_docx
    for path in all_docx:
        errors.extend(inspect_docx(path))

    if errors:
        raise SystemExit("\n".join(errors))
    print(
        f"Final package audit passed: 6 consolidated DOCX + 6 consolidated Markdown; "
        f"{len(final_docx)} section DOCX + {len(final_mds)} mirrored section Markdown; "
        f"anonymous metadata, no comments, no tracked changes, heading styles present."
    )


if __name__ == "__main__":
    main()
