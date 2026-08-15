from __future__ import annotations

import argparse
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


QA_ROOT = Path(__file__).resolve().parent
STAGING = QA_ROOT / "staging"
# This file lives at <project>/docs/v1.0.0/source/ — outputs go to its v1.0.0 parent.
OUTPUT_DIR = QA_ROOT.parent
PROJECT_ROOT = OUTPUT_DIR.parents[1]
MARKDOWN_DIR = OUTPUT_DIR / "markdown"
REFERENCE = QA_ROOT / "_style_reference.docx"
FONT = "Times New Roman"
USABLE_DXA = 9240

DOCS = {
    "SRS_v1_0_0.md": "Software_Requirements_Specification_v1.0.0.docx",
    "NFR_v1_0_0.md": "Non_Functional_Requirements_v1.0.0.docx",
    "UI_UX_Specification_v1_0_0.md": "UI_UX_Specification_v1.0.0.docx",
    "Threat_Model_v1_0_0.md": "Threat_Model_v1.0.0.docx",
    "SDS_v1_0_0.md": "Software_Design_Specification_v1.0.0.docx",
    "DDS_v1_0_0.md": "Database_Design_Specification_v1.0.0.docx",
}


def split_table_row(line: str) -> list[str]:
    text = line.strip().strip("|")
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for char in text:
        if escaped:
            current.append(char)
            escaped = False
        elif char == "\\":
            escaped = True
        elif char == "|":
            cells.append("".join(current).strip().replace("<br>", "\n"))
            current = []
        else:
            current.append(char)
    cells.append("".join(current).strip().replace("<br>", "\n"))
    return cells


def is_control(line: str) -> bool:
    stripped = line.strip()
    return (
        not stripped
        or stripped.startswith("#")
        or stripped.startswith("- ")
        or bool(re.match(r"^\d+\.\s", stripped))
        or stripped.startswith("|")
        or stripped.startswith("<!-- table-widths:")
        or stripped == "---"
        or stripped.startswith("> ")
    )


def markdown_to_blocks(path: Path) -> tuple[str, list[tuple]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError(f"{path.name}: missing title")
    title = lines[0][2:].strip()
    blocks: list[tuple] = []
    pending_widths: list[float] | None = None
    i = 1
    while i < len(lines):
        line = lines[i].strip()
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("<!-- table-widths:"):
            raw = line.removeprefix("<!-- table-widths:").removesuffix("-->").strip()
            pending_widths = [float(v.strip()) for v in raw.split(",") if v.strip()]
            i += 1
            continue
        heading = re.match(r"^(#{2,6})\s+(.+)$", line)
        if heading:
            blocks.append(("heading", len(heading.group(1)), heading.group(2).strip()))
            i += 1
            continue
        if line.startswith("> "):
            blocks.append(("subtitle", line[2:].strip()))
            i += 1
            continue
        if line.startswith("- "):
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:].strip())
                i += 1
            blocks.append(("bullets", items))
            continue
        if re.match(r"^\d+\.\s", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("numbers", items))
            continue
        if line.startswith("|"):
            headers = split_table_row(line)
            if i + 1 >= len(lines) or not lines[i + 1].strip().startswith("|"):
                raise ValueError(f"{path.name}: malformed table near line {i + 1}")
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = split_table_row(lines[i])
                if len(row) != len(headers):
                    raise ValueError(f"{path.name}: table column mismatch near line {i + 1}: {len(row)} != {len(headers)}")
                rows.append(row)
                i += 1
            if pending_widths and len(pending_widths) != len(headers):
                raise ValueError(f"{path.name}: table width count mismatch")
            blocks.append(("table", headers, rows, pending_widths))
            pending_widths = None
            continue
        paragraph = [line]
        i += 1
        while i < len(lines) and not is_control(lines[i]):
            paragraph.append(lines[i].strip())
            i += 1
        blocks.append(("paragraph", " ".join(paragraph)))
    return title, blocks


def normalize_legacy_encoding(text: str) -> str:
    """Repair the limited mojibake sequences inherited from older mirrored Markdown."""
    replacements = {
        "â€“": "–",
        "â€”": "—",
        "â€˜": "‘",
        "â€™": "’",
        "â€œ": "“",
        "â€\u009d": "”",
        "â€¦": "…",
        "Â": "",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text


def set_font(style, size: float, bold: bool = False, italic: bool = False) -> None:
    style.font.name = FONT
    rpr = style._element.get_or_add_rPr()
    for script in ("ascii", "hAnsi", "eastAsia"):
        rpr.rFonts.set(qn(f"w:{script}"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic


def configure_styles(doc: Document) -> None:
    definitions = {
        "Normal": (11, False, False, 3, 5),
        "Title": (24, True, False, 5, 12),
        "Subtitle": (13, False, True, 3, 10),
        "Heading 1": (18, True, False, 12, 5),
        "Heading 2": (15, True, False, 10, 5),
        "Heading 3": (12.5, True, False, 8, 4),
        "Heading 4": (11, True, False, 7, 3),
        "List Bullet": (11, False, False, 1, 2),
        "List Number": (11, False, False, 1, 2),
    }
    for name, values in definitions.items():
        if name not in doc.styles:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[name]
        size, bold, italic, before, after = values
        set_font(style, size, bold, italic)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.03
        if name.startswith("Heading"):
            style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        doc.styles[name].paragraph_format.left_indent = Inches(0.5)
        doc.styles[name].paragraph_format.first_line_indent = Inches(-0.25)


def set_border(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "8")
    border.set(qn("w:space"), "8")
    border.set(qn("w:color"), "A6A6A6")
    pbdr.append(border)
    ppr.append(pbdr)


def add_runs_with_basic_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
        else:
            run = paragraph.add_run(part)


def add_text(doc: Document, text: str, style: str = "Normal"):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runs_with_basic_markdown(p, text)
    return p


def set_table_geometry(table, widths: list[float] | None) -> None:
    count = len(table.columns)
    weights = widths or [1.0] * count
    total = sum(weights)
    values = [round(USABLE_DXA * w / total) for w in weights]
    values[-1] += USABLE_DXA - sum(values)
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tblw.getparent() is None:
        tblpr.append(tblw)
    tblw.set(qn("w:type"), "dxa")
    tblw.set(qn("w:w"), str(USABLE_DXA))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in values:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(value))
        grid.append(node)
    for row in table.rows:
        row._tr.get_or_add_trPr()
        for index, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW")) or OxmlElement("w:tcW")
            if tcw.getparent() is None:
                tcpr.append(tcw)
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(values[index]))


def set_table_borders(table) -> None:
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "A6A6A6")
        borders.append(node)
    tblpr.append(borders)


def add_table(doc: Document, headers, rows, widths) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    hdr = table.rows[0]
    hdrpr = hdr._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    hdrpr.append(repeat)
    for index, value in enumerate(headers):
        cell = hdr.cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E1F2")
        cell._tc.get_or_add_tcPr().append(shd)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(value)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(9.5)
    for row_data in rows:
        row = table.add_row()
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            for part_index, part in enumerate(str(value).split("\n")):
                if part_index:
                    p.add_run().add_break()
                run = p.add_run(part)
                run.font.name = FONT
                run.font.size = Pt(9.5)
    set_table_geometry(table, widths)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    sectpr = body.sectPr
    for child in list(body):
        if child is not sectpr:
            body.remove(child)


def normalize_document(doc: Document, title: str) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        for p in section.header.paragraphs + section.footer.paragraphs:
            p.clear()
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.category = "Product Specification"
    props.subject = "App Lock Version 1.0.0 draft specification"
    props.title = title
    props.keywords = ""
    props.identifier = ""
    props.language = "en"
    props.version = ""


def render_blocks(doc: Document, blocks: list[tuple]) -> None:
    for block in blocks:
        kind = block[0]
        if kind == "heading":
            md_level, text = block[1], block[2]
            word_level = min(max(md_level - 1, 1), 4)
            add_text(doc, text, f"Heading {word_level}")
        elif kind == "subtitle":
            add_text(doc, block[1], "Subtitle")
        elif kind == "paragraph":
            add_text(doc, block[1], "Normal")
        elif kind == "bullets":
            for item in block[1]:
                add_text(doc, item, "List Bullet")
        elif kind == "numbers":
            for item in block[1]:
                add_text(doc, item, "List Number")
        elif kind == "table":
            add_table(doc, block[1], block[2], block[3])
        else:
            raise ValueError(kind)


def build_one(md_path: Path, docx_name: str) -> Path:
    title, blocks = markdown_to_blocks(md_path)
    if not REFERENCE.exists():
        raise FileNotFoundError(REFERENCE)
    output = OUTPUT_DIR / docx_name
    shutil.copyfile(REFERENCE, output)
    doc = Document(output)
    clear_body(doc)
    normalize_document(doc, title)
    configure_styles(doc)
    title_p = add_text(doc, title, "Title")
    set_border(title_p)
    render_blocks(doc, blocks)
    doc.save(output)
    return output


def audit_source(path: Path) -> None:
    title, blocks = markdown_to_blocks(path)
    if "Version 1.0.0" not in path.read_text(encoding="utf-8"):
        raise ValueError(f"{path.name}: version status missing")
    if len(blocks) < 20:
        raise ValueError(f"{path.name}: insufficient content")
    if not title:
        raise ValueError(f"{path.name}: empty title")


def consolidated_source(path: Path) -> Path:
    """Return an ordered SRS source with explicit disposition-only gaps included."""
    if path.name != "SRS_v1_0_0.md":
        return path
    from split_v1_sections import srs_sections

    original = path.read_text(encoding="utf-8").splitlines()
    heading = original[0]
    status = next((x for x in original[1:12] if x.startswith("## Version") or x.startswith("> Version")), "> Version 1.0.0")
    sections = srs_sections(original)
    combined = [heading, "", status, ""]
    for title, body in sections:
        section_body = list(body)
        if title.startswith("4.") and section_body:
            section_body[0] = f"## {title}"
        combined.extend(section_body)
        combined.append("")
    derived = QA_ROOT / "__pycache__" / "SRS_v1_0_0_consolidated.md"
    derived.parent.mkdir(parents=True, exist_ok=True)
    derived.write_text("\n".join(combined).rstrip() + "\n", encoding="utf-8", newline="\n")
    return derived


def build_all() -> None:
    missing = [name for name in DOCS if not (STAGING / name).exists()]
    if missing:
        raise FileNotFoundError(f"Missing staging sources: {missing}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    MARKDOWN_DIR.mkdir(parents=True, exist_ok=True)
    for md_name, docx_name in DOCS.items():
        source = STAGING / md_name
        audit_source(source)
        mirror_name = Path(docx_name).with_suffix(".md").name
        normalized = normalize_legacy_encoding(source.read_text(encoding="utf-8"))
        source.write_text(normalized, encoding="utf-8", newline="\n")
        build_source = consolidated_source(source)
        mirror_text = normalize_legacy_encoding(build_source.read_text(encoding="utf-8"))
        (MARKDOWN_DIR / mirror_name).write_text(mirror_text, encoding="utf-8", newline="\n")
        output = build_one(build_source, docx_name)
        print(f"Built {output}")


def audit_all() -> None:
    for md_name, docx_name in DOCS.items():
        source = STAGING / md_name
        audit_source(source)
        mirror = MARKDOWN_DIR / Path(docx_name).with_suffix(".md").name
        expected = consolidated_source(source)
        if not mirror.exists() or expected.read_bytes() != mirror.read_bytes():
            raise ValueError(f"Mirror mismatch: {mirror}")
        docx = OUTPUT_DIR / docx_name
        if not docx.exists():
            raise FileNotFoundError(docx)
        doc = Document(docx)
        if not doc.paragraphs or doc.paragraphs[0].text.strip() != expected.read_text(encoding="utf-8").splitlines()[0][2:].strip():
            raise ValueError(f"DOCX title mismatch: {docx.name}")
    print("Six-document source, mirror, and package audit passed")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("command", choices=("build", "audit"))
    args = parser.parse_args()
    if args.command == "build":
        build_all()
    else:
        audit_all()


if __name__ == "__main__":
    main()
